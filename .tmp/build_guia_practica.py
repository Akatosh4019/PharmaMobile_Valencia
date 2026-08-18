from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\LENOVO\Desktop\DAM\SESION 2\PROYECTOS\PROJECT2\PharmaMobile")
OUT = ROOT / "output" / "documents" / "Guia_practica_Sesion_2_Roberto_Valencia.docx"
IMG = ROOT / "evidencias" / "guia_practica"

NAVY = RGBColor(20, 48, 74)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(91, 99, 110)
LIGHT = "E8EEF5"


def set_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, size=9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([rpr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("DESARROLLO DE APLICACIONES MÓVILES | SESIÓN 2"), size=8.5, color=GRAY, bold=True)
add_page_number(section.footer.paragraphs[0])

# Editorial cover
doc.add_paragraph().paragraph_format.space_after = Pt(88)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
set_font(p.add_run("GUÍA PRÁCTICA AVANZADA"), size=11, color=BLUE, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
set_font(p.add_run("Kotlin esencial y dominio KMP"), size=28, color=NAVY, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(52)
set_font(p.add_run("Modelado de dominio, null-safety, data classes y pruebas en PharmaMobile"), size=13, color=GRAY, italic=True)

table = doc.add_table(rows=4, cols=2)
table.autofit = False
table.columns[0].width = Inches(1.55)
table.columns[1].width = Inches(4.95)
labels = [
    ("Estudiante", "Roberto Samuel Valencia Saavedra"),
    ("Docente", "Benjamin David Reyna Barreto"),
    ("Proyecto", "PharmaMobile - Kotlin Multiplatform"),
    ("Repositorio", "github.com/Akatosh4019/PharmaMobile_Valencia"),
]
for row, (label, value) in zip(table.rows, labels):
    row.cells[0].width = Inches(1.55)
    row.cells[1].width = Inches(4.95)
    for cell in row.cells:
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(row.cells[0], LIGHT)
    set_font(row.cells[0].paragraphs[0].add_run(label), size=10, color=NAVY, bold=True)
    if label == "Repositorio":
        add_hyperlink(row.cells[1].paragraphs[0], value, "https://github.com/Akatosh4019/PharmaMobile_Valencia")
    else:
        set_font(row.cells[1].paragraphs[0].add_run(value), size=10.5)

doc.add_page_break()

doc.add_heading("1. Propósito y alcance", level=1)
doc.add_paragraph(
    "La práctica consolida el modelado de dominio compartido en commonMain mediante entidades inmutables, "
    "manejo seguro de datos anulables, relaciones entre objetos y validaciones de negocio. Las evidencias "
    "corresponden al proyecto PharmaMobile desarrollado con Kotlin Multiplatform."
)
doc.add_heading("2. Organización del dominio", level=1)
doc.add_paragraph(
    "Los modelos Cliente, Producto, Pedido, DetallePedido y EstadoPedido se encuentran en el paquete "
    "pe.edu.upeu.pharmamobile.domain.model, dentro del código compartido. Esta ubicación permite reutilizar "
    "la lógica de negocio desde Android e iOS."
)
doc.add_heading("3. Evidencias de implementación", level=1)
doc.add_paragraph(
    "Las siguientes capturas documentan la implementación realizada en Android Studio y la validación de una "
    "prueba local de Kotlin, sin depender de un emulador."
)

evidences = [
    ("3.1 Entidad Cliente y null-safety", "01_cliente.png",
     "La entidad Cliente utiliza identificador, nombre y correo obligatorios, mientras que telefono es anulable. "
     "El método obtenertelefono() aplica el operador Elvis para devolver un mensaje seguro cuando no existe un número registrado."),
    ("3.2 Entidad Producto", "02_producto.png",
     "Producto se modela como data class con id, nombre, precio y stock. La inmutabilidad de sus propiedades facilita "
     "el uso de copy() para producir nuevas versiones sin modificar la instancia original."),
    ("3.3 Entidad Pedido y relaciones del dominio", "03_pedido.png",
     "Pedido integra las referencias a Cliente y Producto, además de cantidad, fecha, estado y total. De esta forma se "
     "representa la asociación entre la persona que compra y el producto solicitado."),
    ("3.4 DetallePedido y validaciones", "04_detalle_pedido.png",
     "DetallePedido encapsula la cantidad y el precio unitario de una operación. El bloque init impide cantidades menores "
     "o iguales a cero y precios negativos; el subtotal se obtiene de manera calculada."),
    ("3.5 Estados mediante sealed class", "05_estado_pedido.png",
     "EstadoPedido restringe los estados posibles a Pendiente, Procesando, Entregado y Rechazado. El estado Rechazado "
     "puede transportar el motivo, lo cual permite un tratamiento exhaustivo mediante when."),
    ("3.6 Prueba de Cliente", "06_cliente_test.png",
     "ClienteTest comprueba el comportamiento cuando telefono es null. La ejecución local finaliza con una prueba aprobada "
     "y BUILD SUCCESSFUL, evidenciando que la lógica compartida funciona sin ejecutar el emulador."),
]

for index, (heading, filename, description) in enumerate(evidences):
    if index > 0:
        doc.add_page_break()
    doc.add_heading(heading, level=2)
    p = doc.add_paragraph(description)
    p.paragraph_format.space_after = Pt(8)
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.paragraph_format.keep_with_next = True
    pic.add_run().add_picture(str(IMG / filename), width=Inches(6.45))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(0)
    set_font(cap.add_run(f"Figura {index + 1}. {heading.split(' ', 1)[1]}"), size=9, color=GRAY, italic=True)

doc.add_page_break()
doc.add_heading("4. Resultado de la práctica", level=1)
doc.add_paragraph(
    "Se construyó un dominio compartido con entidades representativas de una farmacia. La solución demuestra el uso de "
    "data classes, tipos anulables, operador Elvis, validaciones con require, propiedades calculadas y una jerarquía "
    "sellada para estados controlados."
)
doc.add_heading("5. Conclusiones", level=1)
conclusions = [
    "El código de dominio permanece en commonMain y puede reutilizarse entre plataformas.",
    "La null-safety evita errores al consultar información opcional del cliente.",
    "Las validaciones de DetallePedido protegen reglas básicas antes de procesar una operación.",
    "La sealed class EstadoPedido delimita los estados válidos y admite información asociada a un rechazo.",
    "La prueba ClienteTest se ejecutó como prueba local y terminó correctamente.",
]
for item in conclusions:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(item)

doc.add_heading("6. Repositorio", level=1)
p = doc.add_paragraph("El código fuente y el historial del proyecto están disponibles en: ")
add_hyperlink(p, "https://github.com/Akatosh4019/PharmaMobile_Valencia", "https://github.com/Akatosh4019/PharmaMobile_Valencia")

doc.core_properties.title = "Guía práctica avanzada - Sesión 2: Kotlin y dominio KMP"
doc.core_properties.subject = "Evidencias de desarrollo del proyecto PharmaMobile"
doc.core_properties.author = "Roberto Samuel Valencia Saavedra"
doc.core_properties.keywords = "Kotlin, KMP, PharmaMobile, dominio, null-safety, pruebas"

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
